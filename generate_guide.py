#!/usr/bin/env python3
"""Generate the SF Config Debt Scanner Integration Guide as a Word document."""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ---------- Cover Page ----------
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SAP SuccessFactors\nAI Agent Integration Guide")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Config Debt Scanner via MCP\n(Model Context Protocol)")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("For SAP SuccessFactors Consultants & Administrators")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ---------- TOC ----------
doc.add_heading("Contents", level=1)
toc_items = [
    "What is an AI Agent? A Simple Explanation",
    "What is MCP? (Model Context Protocol)",
    "Why This Matters for SAP SuccessFactors",
    "The Config Debt Scanner - What It Does",
    "Who Can Use This? (Supported AI Agents)",
    "Integration Guide for Each Platform",
    "Example Conversations With the Scanner",
    "Privacy & Data Security",
    "FAQ",
    "Getting Help",
]
for item in toc_items:
    p = doc.add_paragraph(item, style="List Number")
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ---------- Section 1 ----------
doc.add_heading("What is an AI Agent? A Simple Explanation", level=1)

doc.add_paragraph(
    "An AI agent is like a very smart assistant that can understand what you ask and then "
    "take action - not just give you an answer. Think of it as having a junior consultant "
    "who knows everything about SAP SuccessFactors configuration and can run checks, "
    "analyse data, and produce reports whenever you ask."
)

doc.add_paragraph(
    "Unlike a search engine or chatbot that just returns text, an AI agent can:"
)
for b in [
    "Connect to systems (like your SF tenant via OData)",
    "Read files and extract insights",
    "Run analyses and return structured results",
    "Explain findings in plain English",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph(
    "You interact with it through a chat window - like WhatsApp or Teams - but instead of "
    "talking to a person, you are telling an AI what to investigate."
)

# ---------- Section 2 ----------
doc.add_heading("What is MCP? (Model Context Protocol)", level=1)

doc.add_paragraph(
    "MCP stands for Model Context Protocol. It is a standard way for AI agents to connect "
    "to tools and data sources - like a USB plug for AI."
)

doc.add_paragraph(
    "Before MCP, every AI tool had its own custom way of connecting. This meant each "
    "integration had to be built from scratch, and switching AI assistants meant rebuilding "
    "everything."
)

doc.add_paragraph(
    "With MCP, a tool (like our Config Debt Scanner) publishes a list of things it can do. "
    "Any AI agent that speaks MCP can discover those capabilities automatically and use them. "
    "This is why the same Config Debt Scanner works with Hermes Agent, Claude Code, Cursor, "
    "and others - no custom coding for each platform."
)

p = doc.add_paragraph()
run = p.add_run("Key point: ")
run.bold = True
p.add_run("If you know how to use one MCP tool, you know how to use them all. "
           "The integration process is identical for every MCP server.")

# ---------- Section 3 ----------
doc.add_heading("Why This Matters for SAP SuccessFactors", level=1)

doc.add_paragraph(
    "SAP SuccessFactors tenants accumulate configuration debt over time - custom fields, "
    "MDF objects, business rules, event reasons, picklists. Most consultants track this "
    "manually or through periodic audits."
)

doc.add_paragraph("With this AI Agent integration, you can:")
for b in [
    'Ask in plain English: "Scan our tenant for custom field sprawl" - and get an answer in seconds',
    "Run audits without logging into Provisioning or OData explorers",
    "Share findings with clients in structured reports",
    "Generate assessment questions for client workshops on demand",
    "Automate recurring checks (e.g., every release cycle)",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph(
    "The barrier is intentionally low. You do not need to write code. You do not need to "
    "understand APIs. The AI agent handles all of that."
)

# ---------- Section 4 ----------
doc.add_heading("The Config Debt Scanner - What It Does", level=1)

doc.add_paragraph(
    "The SF Config Debt Scanner is a tool that analyses an SAP SuccessFactors tenant and "
    "identifies configuration debt. It connects via the standard OData v2 API and reads "
    "only metadata and record counts - not employee data."
)

doc.add_paragraph("It can do the following tasks, accessible through any MCP-compatible AI agent:")

table = doc.add_table(rows=8, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["Tool / Task", "What It Does"]):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

tools_data = [
    ("Scan Metadata XML", "Analyse raw $metadata XML for config debt indicators - no tenant connection needed. Upload or paste your metadata XML and get a full report."),
    ("Test Connection", "Check if an SF tenant is reachable via OData v2. Returns connectivity status and record counts for core entities (EmpJob, User, Position, etc.)."),
    ("Full Tenant Scan", "Connect to a live SF tenant, pull the $metadata, run schema analysis and Tier 1 count checks, and return a config debt score with a 90-day remediation roadmap."),
    ("Assessment Questions", "Generate guided configuration debt assessment questions across 7 categories: governance, custom fields, MDF, picklists, event reasons, foundation objects, and business rules."),
    ("Rate Findings", "Score a set of custom findings and get an overall debt score with area breakdown and priority-ranked next steps."),
    ("Known Entities", "List the EC entity names the scanner recognises for classification purposes."),
    ("About", "Get server version, available tools list, and data policy information."),
]
for row_idx, (tool, desc) in enumerate(tools_data, start=1):
    table.rows[row_idx].cells[0].text = tool
    table.rows[row_idx].cells[1].text = desc

doc.add_paragraph()

# ---------- Section 5 ----------
doc.add_heading("Who Can Use This? (Supported AI Agents)", level=1)

doc.add_paragraph(
    "The Config Debt Scanner works with any AI agent that supports the MCP protocol. "
    "The most common ones for SAP consultants are:"
)

platforms = [
    ("Hermes Agent", "A terminal-based AI assistant by Nous Research. Best for technical users comfortable with a command line. Free and open-source."),
    ("Claude Code", "Anthropic's AI coding assistant that runs in the terminal. Available via subscription."),
    ("Cursor", "An AI code editor with built-in MCP support. More visual than terminal-based agents."),
    ("VS Code + Extensions", "Microsoft's code editor with MCP plugin support. Familiar interface for those already using VS Code for XML/JSON editing."),
]

for i, (name, desc) in enumerate(platforms, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {name}")
    run.bold = True
    p.add_run(f" - {desc}")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Note: ")
run.bold = True
p.add_run("The setup for each platform takes approximately 2-5 minutes and is a one-time task.")

# ---------- Section 6 ----------
doc.add_heading("Integration Guide", level=1)

doc.add_paragraph(
    "Below are the exact steps to connect the Config Debt Scanner to each AI agent. "
    "These are one-time setup steps. Once done, the tools are available every time you use the AI agent."
)

# Hermes
doc.add_heading("Hermes Agent", level=2)
doc.add_paragraph("1. Install Hermes Agent (see hermes-agent.nousresearch.com/download)")
doc.add_paragraph("2. Ensure you have Python 3.10+ and the scanner code downloaded")
doc.add_paragraph("3. Install dependencies:")
p = doc.add_paragraph()
run = p.add_run("   pip install -r requirements.txt")
run.font.name = "Consolas"
run.font.size = Pt(9)
doc.add_paragraph("4. Add these lines to your ~/.hermes/config.yaml:")
code_block = (
    "mcp_servers:\n"
    "  sf-config-debt-scanner:\n"
    '    command: "/path/to/run_mcp_server.sh"\n'
    "    enabled: true\n"
)
p = doc.add_paragraph()
run = p.add_run(code_block)
run.font.name = "Consolas"
run.font.size = Pt(9)
doc.add_paragraph("5. Restart Hermes. The 7 tools will auto-discover.")

# Claude Code
doc.add_heading("Claude Code (by Anthropic)", level=2)
doc.add_paragraph("1. Ensure Claude Code is installed (requires Anthropic subscription)")
doc.add_paragraph("2. Add to ~/.claude/claude_desktop_config.json:")
code_block2 = (
    '{\n'
    '  "mcpServers": {\n'
    '    "sf-config-debt-scanner": {\n'
    '      "command": "python",\n'
    '      "args": ["/full/path/to/mcp_server.py"]\n'
    '    }\n'
    '  }\n'
    '}\n'
)
p = doc.add_paragraph()
run = p.add_run(code_block2)
run.font.name = "Consolas"
run.font.size = Pt(9)
doc.add_paragraph("3. Restart Claude Code. The tools will appear automatically.")

# Cursor
doc.add_heading("Cursor", level=2)
for s in [
    'Open Cursor Settings > Features > MCP Servers',
    'Click "Add New MCP Server"',
    'Set Name to: sf-config-debt-scanner',
    'Set Type to: command',
    'Set Command to: python /full/path/to/mcp_server.py',
    'Click Save. The tools will appear in the MCP panel.',
]:
    doc.add_paragraph(s)

# ---------- Section 7 ----------
doc.add_heading("Example Conversations", level=1)

doc.add_paragraph(
    "Once connected, you can talk to the AI agent naturally. Here are examples of "
    "what you can say:"
)

examples = [
    ("Metadata scan (no tenant needed)",
     'You: "I have the $metadata XML from a tenant. Can you scan it for configuration debt?"\n\n'
     "[Paste the XML content]\n\n"
     "Agent: Shows entity breakdown, custom field concentration, field count warnings, "
     "debt score, and remediation roadmap."),
    ("Tenant connection test",
     'You: "Can you test the connection to api55.sapsf.eu using basic auth?"\n\n'
     'Agent: "Yes. Please provide your username and password."\n'
     "[Provide credentials]\n\n"
     "Agent: Confirms connectivity and shows record counts for EmpJob, User, Position, etc."),
    ("Full tenant scan",
     'You: "Run a full config debt scan against our production tenant."\n\n'
     "Agent: Guides you through providing the connection details.\n\n"
     "Agent: Returns a full report with debt score (e.g., 42/100 - Moderate), "
     "categorized findings, and a 90-day action plan."),
    ("Workshop preparation",
     'You: "Give me assessment questions for governance and custom fields."\n\n'
     "Agent: Returns structured assessment questions with severity ratings and "
     "rationale for each, ready to use in your workshop."),
    ("Custom finding scoring",
     'You: "Here are 15 findings from our audit. Score these and give me a roadmap."\n\n'
     "Agent: Scores each finding, calculates overall debt score, and produces "
     "a priority-ranked 90-day roadmap."),
]

for title, content in examples:
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

# ---------- Section 8 ----------
doc.add_heading("Privacy & Data Security", level=1)

privacy_points = [
    "Zero employee data is ever stored by the scanner. It reads only metadata (entity definitions, field names, relationships) and record counts.",
    "No personally identifiable information (PII), payroll values, or individual records are accessed or retained.",
    "The scanner connects via standard OData v2 API - the same API used by SF reporting and integration tools.",
    "Authentication credentials are provided per-session and are never saved by the scanner.",
    "All processing happens locally within your environment. No data is sent to external servers.",
    "The scanner is fully self-contained Python code that can be reviewed and audited before use.",
]
for pt in privacy_points:
    doc.add_paragraph(pt, style="List Bullet")

# ---------- Section 9 ----------
doc.add_heading("FAQ", level=1)

faqs = [
    ("Do I need coding skills to use this?",
     "No. The AI agent handles all the technical interaction. You communicate in plain English. "
     "The one-time setup requires someone technical to configure the MCP plugin, but daily use "
     "is purely conversational."),
    ("Does this connect to my live production tenant?",
     "Yes, it connects via OData v2 using the same API endpoints your existing tools use. "
     "It reads schema metadata and record counts only - never employee data."),
    ("Is this SAP-approved?",
     "The scanner uses standard SAP SuccessFactors OData v2 APIs. It operates as a read-only "
     "analytics tool and does not modify any tenant configuration."),
    ("Can multiple people use the same scanner?",
     "Yes. Once the MCP server is configured on a shared machine or server, anyone with access "
     "to the AI agent can use it."),
    ("How often should I scan?",
     "Quarterly scans are recommended for most tenants, or per-release-cycle (every 6 months "
     "coinciding with SAP SF releases). The 90-day roadmap helps prioritize which findings to "
     "address first."),
    ("What if I don't have an AI agent installed?",
     "You can still use the scanner directly via the Python command line or through the browser "
     "dashboard. See the README file for CLI usage."),
    ("Is there a cost?",
     "The scanner itself is free and open-source. The AI agent you use to interact with it may "
     "have its own subscription (e.g., Claude Code) or be free (e.g., Hermes Agent)."),
]

for q, a in faqs:
    doc.add_heading(q, level=2)
    doc.add_paragraph(a)

# ---------- Section 10 ----------
doc.add_heading("Getting Help", level=1)

doc.add_paragraph(
    "For questions about the Config Debt Scanner, refer to the project repository:\n"
    "https://github.com/SahirVhora/sf-config-debt-radar\n\n"
    "For questions about the AI agent platform:\n"
    "- Hermes Agent: hermes-agent.nousresearch.com/docs\n"
    "- Claude Code: docs.anthropic.com/en/docs/claude-code/\n"
    "- Cursor: docs.cursor.com/get-started/mcp\n"
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Document generated June 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# Save
output_path = os.path.expanduser("~/projects/sapsf/sf-config-debt-radar/SF_Config_Debt_Scanner_Integration_Guide.docx")
doc.save(output_path)
print(f"Saved to {output_path}")
